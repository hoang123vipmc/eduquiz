import json
import sys
from typing import List, Dict, Any

# Fix for Windows console unicode printing
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
except ImportError:
    print("Thư viện 'python-docx' chưa được cài đặt.")
    print("Vui lòng chạy lệnh: pip install -r requirements.txt")
    sys.exit(1)

def load_json_data(file_path: str) -> Dict[str, Any]:
    """
    Đọc dữ liệu từ file JSON và trả về dictionary.
    Xử lý các ngoại lệ khi file không tồn tại hoặc lỗi định dạng.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except FileNotFoundError:
        print(f"Lỗi: Không tìm thấy file '{file_path}'. Vui lòng kiểm tra lại đường dẫn.")
        return {}
    except json.JSONDecodeError:
        print(f"Lỗi: File '{file_path}' không phải định dạng JSON hợp lệ.")
        return {}
    except Exception as e:
        print(f"Đã xảy ra lỗi không xác định khi đọc file: {e}")
        return {}

def extract_questions(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    """
    Trích xuất danh sách câu hỏi, các lựa chọn và đáp án đúng.
    """
    extracted_questions = []
    
    try:
        # Tùy thuộc vào cấu trúc JSON, lấy danh sách các phần (sections/data)
        sections = data.get('data', [])
        question_counter = 1
        
        for section in sections:
            questions = section.get('questions', [])
            for q in questions:
                # Trích xuất nội dung câu hỏi
                question_text = q.get('name', 'Không có nội dung câu hỏi')
                answers = q.get('answers', [])
                
                parsed_answers = []
                correct_answer = "Không có đáp án đúng"
                
                # Mã ASCII cho 'A' là 65, dùng để đánh dấu A, B, C, D...
                char_code = 65 
                
                for ans in answers:
                    option_text = ans.get('option', '')
                    is_correct = ans.get('is_correct', False)
                    letter = chr(char_code)
                    
                    option_formatted = f"{letter}. {option_text}"
                    parsed_answers.append(option_formatted)
                    
                    if is_correct:
                        correct_answer = option_formatted
                        
                    char_code += 1
                
                extracted_questions.append({
                    'index': question_counter,
                    'text': question_text,
                    'options': parsed_answers,
                    'correct': correct_answer
                })
                question_counter += 1
                
    except Exception as e:
        print(f"Lỗi trong quá trình trích xuất dữ liệu: {e}")
        
    return extracted_questions

def export_to_txt(questions: List[Dict[str, Any]], output_path: str):
    """
    Xuất danh sách câu hỏi ra file TXT với format đẹp.
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            for q in questions:
                f.write(f"Câu {q['index']}: {q['text']}\n")
                for opt in q['options']:
                    f.write(f"{opt}\n")
                f.write(f"=> Đáp án đúng: {q['correct']}\n")
                f.write("-" * 50 + "\n\n")
        print(f"Đã xuất file TXT thành công tại: {output_path}")
    except Exception as e:
        print(f"Lỗi khi xuất file TXT: {e}")

def export_to_docx(questions: List[Dict[str, Any]], output_path: str):
    """
    Xuất danh sách câu hỏi ra file Word (DOCX) với format đẹp.
    """
    try:
        doc = Document()
        doc.add_heading('Đề Cương Ôn Tập Trắc Nghiệm', 0)
        
        for q in questions:
            # Nội dung câu hỏi (in đậm)
            p_question = doc.add_paragraph()
            run_question = p_question.add_run(f"Câu {q['index']}: {q['text']}")
            run_question.bold = True
            
            # Các lựa chọn
            for opt in q['options']:
                doc.add_paragraph(opt)
                
            # Đáp án đúng (in nghiêng)
            p_correct = doc.add_paragraph()
            run_correct = p_correct.add_run(f"=> Đáp án đúng: {q['correct']}")
            run_correct.italic = True
            
            # Thêm khoảng trắng giữa các câu
            doc.add_paragraph()
            
        doc.save(output_path)
        print(f"Đã xuất file DOCX thành công tại: {output_path}")
    except Exception as e:
        print(f"Lỗi khi xuất file DOCX: {e}")

def main():
    input_file = 'sample_data.json'
    txt_output = 'de_cuong.txt'
    docx_output = 'de_cuong.docx'
    
    print("="*50)
    print("CHƯƠNG TRÌNH TRÍCH XUẤT DỮ LIỆU TRẮC NGHIỆM")
    print("="*50)
    
    # 1. Đọc dữ liệu
    data = load_json_data(input_file)
    if not data:
        print("Không có dữ liệu để xử lý. Kết thúc chương trình.")
        return
        
    # 2. Trích xuất câu hỏi
    questions = extract_questions(data)
    print(f"Đã trích xuất thành công {len(questions)} câu hỏi.\n")
    
    # 3. Xuất file
    if questions:
        export_to_txt(questions, txt_output)
        export_to_docx(questions, docx_output)
        print("\nHoàn tất quá trình trích xuất!")
    else:
        print("Không tìm thấy câu hỏi nào trong dữ liệu JSON.")

if __name__ == "__main__":
    main()
